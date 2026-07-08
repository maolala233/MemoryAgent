"""Document parsing for PDF / DOCX / Markdown / TXT uploads."""
from __future__ import annotations

import hashlib
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from ..utils.logger import warn


class DocumentParser:
    """Extracts text and metadata from common document formats."""

    def parse(self, file_path: Path) -> Dict[str, Any]:
        suffix = file_path.suffix.lower()
        if suffix == ".md":
            return self.parse_markdown(file_path)
        if suffix == ".txt":
            return self.parse_text(file_path)
        if suffix == ".pdf":
            return self.parse_pdf(file_path)
        if suffix == ".docx":
            return self.parse_docx(file_path)
        # Fallback: treat as text
        return self.parse_text(file_path)

    def parse_text(self, file_path: Path) -> Dict[str, Any]:
        text = file_path.read_text(encoding="utf-8", errors="ignore")
        return {
            "text": text,
            "metadata": {"filename": file_path.name, "format": "txt"},
            "pages": 1,
        }

    def parse_markdown(self, file_path: Path) -> Dict[str, Any]:
        text = file_path.read_text(encoding="utf-8", errors="ignore")
        sections = self._extract_md_sections(text)
        return {
            "text": text,
            "metadata": {
                "filename": file_path.name,
                "format": "md",
                "sections": [s[0] for s in sections],
            },
            "pages": 1,
            "sections": sections,
        }

    def parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        try:
            import pypdf  # type: ignore
        except ImportError:
            try:
                import PyPDF2 as pypdf  # type: ignore
            except ImportError:
                warn("pypdf/PyPDF2 not installed; reading PDF as binary")
                return {"text": "", "metadata": {"filename": file_path.name, "format": "pdf"},
                        "pages": 0}
        text_parts: List[str] = []
        try:
            reader = pypdf.PdfReader(str(file_path))
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
        except Exception as exc:
            warn(f"PDF parse failed: {exc}")
        text = "\n\n".join(text_parts)
        return {
            "text": text,
            "metadata": {"filename": file_path.name, "format": "pdf"},
            "pages": len(text_parts),
        }

    def parse_docx(self, file_path: Path) -> Dict[str, Any]:
        try:
            import docx  # type: ignore
        except ImportError:
            warn("python-docx not installed; reading docx as binary")
            return {"text": "", "metadata": {"filename": file_path.name, "format": "docx"},
                    "pages": 0}
        doc = docx.Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 遍历文档 body 中的所有元素，按出现顺序混合段落与表格
        from docx.oxml.ns import qn
        body = doc.element.body
        ordered_chunks: list[str] = []
        para_iter = iter(doc.paragraphs)
        tbl_iter = iter(doc.tables)
        for child in body.iterchildren():
            tag = child.tag
            if tag == qn("w:p"):
                try:
                    p = next(para_iter)
                    t = p.text.strip()
                    if t:
                        ordered_chunks.append(t)
                except StopIteration:
                    pass
            elif tag == qn("w:tbl"):
                try:
                    tbl = next(tbl_iter)
                    for row in tbl.rows:
                        cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                        if cells:
                            ordered_chunks.append(" | ".join(cells))
                except StopIteration:
                    pass
        if not ordered_chunks:
            ordered_chunks = paragraphs
        text = "\n\n".join(ordered_chunks)
        return {
            "text": text,
            "metadata": {"filename": file_path.name, "format": "docx",
                         "paragraph_count": len(paragraphs),
                         "table_count": len(doc.tables)},
            "pages": 1,
        }

    def extract_metadata(self, content: str, filename: str) -> Dict[str, Any]:
        title = filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " ").title()
        for line in content.splitlines():
            line = line.strip()
            if line.startswith("# "):
                title = line[2:].strip()
                break
        word_count = len(content.split())
        return {
            "title": title,
            "filename": filename,
            "word_count": word_count,
            "char_count": len(content),
        }

    @staticmethod
    def _extract_md_sections(text: str) -> List[Tuple[str, str]]:
        sections: List[Tuple[str, str]] = []
        current_title = "Introduction"
        buffer: List[str] = []
        for line in text.splitlines():
            if line.startswith("#"):
                if buffer:
                    sections.append((current_title, "\n".join(buffer).strip()))
                    buffer = []
                current_title = line.lstrip("#").strip() or current_title
            else:
                buffer.append(line)
        if buffer:
            sections.append((current_title, "\n".join(buffer).strip()))
        return sections


document_parser = DocumentParser()
